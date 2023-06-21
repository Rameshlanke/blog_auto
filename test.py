import httplib2
from oauth2client.client import flow_from_clientsecrets
from oauth2client.file import Storage
from oauth2client.tools import run_flow
from googleapiclient import discovery
import mammoth

import os
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor

url = "https://answersq.com/udemy-paid-courses-for-free-with-certificate/"
response = requests.get(url)
html_content = response.content

soup = BeautifulSoup(html_content, 'html.parser')
ul_elements = soup.find_all('ul')
document = Document()
for ul in ul_elements:
   
    document.add_section()
    for li in ul.find_all('li'):
        p = document.add_paragraph(style='List Bullet')

        p.add_run(li.text.strip())

        # Check if the li element contains an href link
        if li.a:
            # Add a space and the href link to the paragraph
            p.add_run(f' ({li.a["href"]})').italic = True

document.save('post.docx')
print("Created post.docx")


from docx import Document
from docxtpl import DocxTemplate
source_doc = Document('source.docx')

# Load the destination Word document
destination_doc = Document('post.docx')

# Create a new, empty document for the modified document
modified_doc = Document()

# Get the text of the source document
source_text = '\n'.join([para.text for para in source_doc.paragraphs])

for para in destination_doc.paragraphs:
    # Remove the string "- Enroll for Free" from the paragraph's text
    para_text = para.text.replace('– Enroll for Free', '').strip()
    
    if para_text and para_text not in source_text:
        modified_para = modified_doc.add_paragraph(para_text)

# Save the modified document to a new file
modified_doc.save('modified.docx')
print("Created modified.docx")




#code in ................

import os
import win32com.client as win32

# Get current working directory
cwd = os.getcwd()

# Create instance of Word application
word = win32.gencache.EnsureDispatch('Word.Application')

# Open the document
doc = word.Documents.Open(os.path.join(cwd, 'modified.docx'))

# Apply autoformatting
doc.Content.AutoFormat()

# Save the document
doc.Save()

# Close the document and quit Word application
doc.Close()
word.Quit()
print("Autoformatting Finished")

#code out.............




import docx

# Get the path to the word document file
document_path = "modified.docx"

# Open the word document file
document = docx.Document(document_path)

# Get the paragraphs of the word document file
paragraphs = document.paragraphs

# Search for the string in the text
search_string = "answersq"

# Check if the string is found
found = False
for paragraph in paragraphs:
    # print(paragraph)
    if search_string in paragraph.text:
        found = True
        break

if found:
    print("found and Stopped the program")
else:


    # Start the OAuth flow to retrieve credentials
    def authorize_credentials():
        CLIENT_SECRET = 'client_secret.json'
        SCOPE = 'https://www.googleapis.com/auth/blogger'
        STORAGE = Storage('credentials.storage')
        # Fetch credentials from storage
        credentials = STORAGE.get()
        # If the credentials doesn't exist in the storage location then run the flow
        if credentials is None or credentials.invalid:
            flow = flow_from_clientsecrets(CLIENT_SECRET, scope=SCOPE)
            http = httplib2.Http()
            credentials = run_flow(flow, STORAGE, http=http)
        return credentials

    # print(credentials)
    def getBloggerService():
        credentials = authorize_credentials()
        http = credentials.authorize(httplib2.Http())
        discoveryUrl = ('https://blogger.googleapis.com/$discovery/rest?version=v3')
        service = discovery.build('blogger', 'v3', http=http, discoveryServiceUrl=discoveryUrl)
        return service

    def postToBlogger(payload):
        service = getBloggerService()
        post=service.posts()
        insert=post.insert(blogId='5651255862571344247',body=payload).execute()
        print("Done post!")
        return insert


    #IHTREEK TECH COURSES BLOG ID: 5651255862571344247
    #Test blog ID: 3733272535113126568



    import mammoth

    def buildHtml(file_path):
        if file_path.endswith('.html'):
            with open(file_path, 'r') as file:
                html = file.read()
        elif file_path.endswith('.docx'):
            with open(file_path, 'rb') as file:
                result = mammoth.convert_to_html(file)
                html = result.value
        else:
            raise ValueError('Invalid file extension. Only .html and .docx files are supported.')

        return html
    html1 = buildHtml('first.html')
    html2 = buildHtml('modified.docx')
    html3 = buildHtml('file2.docx')
    all_html = html1 + html2  + html3

    from datetime import date

    # Get the current date in the format YYYY-MM-DD
    today = date.today().strftime("%d-JUNE-%y")

    # Set the title to the current date
    # title = input()
    title = f"UDEMY COURSES FREE CERTIFICATE | {today} | IHTREEKTECHCOURSES"
    posttoday = date.today().strftime("JUNE %d")

    print(title)

    # Output: "Post for 2023-04-19"
    # title = "Testing post 2" 
    # print(htmlData)
    permalink = f"https://www.ihtreektechcourses.com/udemy-courses-free-certificates-{posttoday}-ihtreektech"
    search_description = "Free courses from Udemy to help you make the most of your time, from working at home to trending technical skills and self-improvement, wherever you are."

    customMetaData = "This is meta data"
    payload={
            "content": all_html,
            "title": title,
            'labels': ['May 2023','Udemy Free','IHTREEKTECHCOURSES'],
            'customMetaData': customMetaData,
            'url': permalink,
            'Description': search_description
        }
    postToBlogger(payload)

    import os

    # List of file names to delete
    file_list = ['post.docx','modified.docx']
    # Iterate over the file names and delete each file
    for file in file_list:
        os.remove(file)
        print(f"{file} deleted successfully!")


